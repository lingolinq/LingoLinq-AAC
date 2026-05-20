#!/usr/bin/env python3
"""LingoLinq-branded .docx of the 1Password decision memo.

Tier 2 Internal, Working shape per BRAND.md v3.
Muted colorway, Calibri throughout, Scot's signatures applied.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

DEEP_PURPLE = RGBColor(0x22, 0x21, 0x5A)
PERIWINKLE = RGBColor(0x78, 0x86, 0xF3)
DARK_SLATE = RGBColor(0x1C, 0x1C, 0x2D)
DUSTY_PERIWINKLE = RGBColor(0x89, 0x93, 0xC0)
CHARCOAL = RGBColor(0x37, 0x41, 0x51)
AMBER_GOLD = RGBColor(0xF2, 0xB9, 0x5A)
AMBER_LIGHT_HEX = "FFFBF0"
AMBER_DARK = RGBColor(0x85, 0x4F, 0x0B)
LIGHT_GRAY = RGBColor(0xF0, 0xF0, 0xF0)
OFF_WHITE_HEX = "F7F7FB"
PURPLE_LIGHT_HEX = "EEEDFE"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TEAL = RGBColor(0x2A, 0x9D, 0x8F)
RED = RGBColor(0xCC, 0x00, 0x00)

BODY_FONT = "Calibri"
LOGO_PATH = "/mnt/c/Users/scotw/Projects/LingoLinq-Brand/logos/PNG/Muted/LL_Line_Muted.png"

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(11)
normal.font.color.rgb = CHARCOAL
normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
normal.paragraph_format.line_spacing = 1.4
normal.paragraph_format.space_after = Pt(6)


def set_cell_shade(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color="8993C0", size="4", left=False, right=False, top=False, bottom=False, all_sides=False):
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcBorders"))
    if existing is not None:
        tc_pr.remove(existing)
    tc_borders = OxmlElement("w:tcBorders")
    sides = []
    if all_sides:
        sides = ["top", "left", "bottom", "right"]
    else:
        if top:
            sides.append("top")
        if left:
            sides.append("left")
        if bottom:
            sides.append("bottom")
        if right:
            sides.append("right")
    for side in sides:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def remove_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcBorders"))
    if existing is not None:
        tc_pr.remove(existing)
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def remove_table_borders(table):
    tbl_pr = table._element.tblPr
    if tbl_pr is None:
        return
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    tbl_pr.append(borders)


def style_run(run, *, size=11, color=CHARCOAL, bold=False, italic=False, font=BODY_FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_paragraph(text="", *, size=11, color=CHARCOAL, bold=False, italic=False,
                  align=None, space_before=0, space_after=6, line_spacing=1.4,
                  parent=None):
    p = parent.add_paragraph() if parent is not None else doc.add_paragraph()
    if text:
        run = p.add_run(text)
        style_run(run, size=size, color=color, bold=bold, italic=italic)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = line_spacing
    return p


def add_runs(runs, *, parent=None, space_before=0, space_after=6, line_spacing=1.4, align=None):
    p = parent.add_paragraph() if parent is not None else doc.add_paragraph()
    for chunk, opts in runs:
        run = p.add_run(chunk)
        style_run(
            run,
            size=opts.get("size", 11),
            color=opts.get("color", CHARCOAL),
            bold=opts.get("bold", False),
            italic=opts.get("italic", False),
        )
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = line_spacing
    return p


def add_h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, size=18, color=DARK_SLATE, bold=True)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, size=13, color=DUSTY_PERIWINKLE, bold=True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_italic_intro(text):
    return add_paragraph(text, italic=True, color=CHARCOAL, space_after=6)


def add_bullet(text, *, bold_lead=None, parent=None):
    p = parent.add_paragraph(style="List Bullet") if parent is not None else doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r1 = p.add_run(bold_lead)
        style_run(r1, bold=True, color=DARK_SLATE)
        r2 = p.add_run(text)
        style_run(r2)
    else:
        run = p.add_run(text)
        style_run(run)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.35
    return p


def add_periwinkle_divider(parent=None, color="8993C0", thickness="6"):
    p = (parent.add_paragraph() if parent else doc.add_paragraph())
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), thickness)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)


def header_strip():
    """Compact logo header on page 1: logo left, title block right, periwinkle bottom border."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(table)

    table.columns[0].width = Inches(2.6)
    table.columns[1].width = Inches(4.0)

    left = table.rows[0].cells[0]
    left.width = Inches(2.6)
    remove_cell_borders(left)
    left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    left_p = left.paragraphs[0]
    left_p.paragraph_format.space_after = Pt(0)
    if os.path.exists(LOGO_PATH):
        run = left_p.add_run()
        run.add_picture(LOGO_PATH, width=Inches(2.2))
    else:
        r = left_p.add_run("LingoLinq")
        style_run(r, size=18, color=DARK_SLATE, bold=True)

    right = table.rows[0].cells[1]
    right.width = Inches(4.0)
    remove_cell_borders(right)
    right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    title_p = right.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title_p.paragraph_format.space_after = Pt(0)
    title_run = title_p.add_run("1Password Decision")
    style_run(title_run, size=20, color=DARK_SLATE, bold=True)

    sub_p = right.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sub_p.paragraph_format.space_after = Pt(0)
    sub_run = sub_p.add_run("Internal Memo  |  2026-04-17")
    style_run(sub_run, size=10, color=DUSTY_PERIWINKLE, italic=True)

    add_periwinkle_divider()


def title_banner():
    """Dark-filled title banner: Deep Purple bg, white text. Report name + period + byline."""
    table = doc.add_table(rows=1, cols=1)
    remove_table_borders(table)
    cell = table.rows[0].cells[0]
    set_cell_shade(cell, "22215A")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    line1 = cell.paragraphs[0]
    line1.paragraph_format.space_after = Pt(2)
    r1 = line1.add_run("1Password Decision")
    style_run(r1, size=18, color=WHITE, bold=True)

    line2 = cell.add_paragraph()
    line2.paragraph_format.space_after = Pt(0)
    r2 = line2.add_run("$23.97/month  |  Sign-off needed by 2026-04-18  |  Scot Wahlquist")
    style_run(r2, size=11, color=WHITE, italic=True)

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side in ("top", "bottom", "left", "right"):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), "200" if side in ("left", "right") else "140")
        m.set(qn("w:type"), "dxa")
        tc_mar.append(m)
    tc_pr.append(tc_mar)

    add_paragraph("", space_after=4)


def amber_callout(label, body_text):
    """Signature 1: Bottom-line opener in amber callout with left bar."""
    table = doc.add_table(rows=1, cols=1)
    remove_table_borders(table)
    cell = table.rows[0].cells[0]
    set_cell_shade(cell, AMBER_LIGHT_HEX)
    set_cell_borders(cell, color="F2B95A", size="24", left=True)

    label_p = cell.paragraphs[0]
    label_p.paragraph_format.space_after = Pt(2)
    label_run = label_p.add_run(label)
    style_run(label_run, size=10, color=AMBER_DARK, bold=True)

    body_p = cell.add_paragraph()
    body_p.paragraph_format.space_after = Pt(0)
    body_run = body_p.add_run(body_text)
    style_run(body_run, size=11, color=CHARCOAL)

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, val in (("top", "120"), ("bottom", "120"), ("left", "180"), ("right", "180")):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), val)
        m.set(qn("w:type"), "dxa")
        tc_mar.append(m)
    tc_pr.append(tc_mar)

    add_paragraph("", space_after=4)


def info_callout(title_text, paragraphs):
    """Soft info callout: Purple Light fill + Dusty Periwinkle left bar."""
    table = doc.add_table(rows=1, cols=1)
    remove_table_borders(table)
    cell = table.rows[0].cells[0]
    set_cell_shade(cell, PURPLE_LIGHT_HEX)
    set_cell_borders(cell, color="8993C0", size="20", left=True)

    title_p = cell.paragraphs[0]
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run(title_text)
    style_run(title_run, size=11, color=DARK_SLATE, bold=True)

    for para in paragraphs:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(para)
        style_run(run, color=CHARCOAL)

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, val in (("top", "140"), ("bottom", "140"), ("left", "180"), ("right", "180")):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), val)
        m.set(qn("w:type"), "dxa")
        tc_mar.append(m)
    tc_pr.append(tc_mar)

    add_paragraph("", space_after=4)


def cost_table():
    """Dark Slate header row, off-white alternating rows, no harsh borders."""
    headers = ["Option", "What it covers", "Cost for us"]
    rows_data = [
        ("Bitwarden Free", "Personal use, one user only", "$0, cannot be used for business"),
        ("Bitwarden Teams ($4/user)", "Small teams, no BAA", "Not HIPAA-eligible, ruled out"),
        ("Bitwarden Enterprise ($6/user)", "Comparable to 1Password Business", "$18/mo for 3 seats"),
        ("1Password Business ($7.99/user)", "What we have now", "$24/mo for 3 seats"),
    ]
    table = doc.add_table(rows=len(rows_data) + 1, cols=3)
    remove_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shade(cell, "1C1C2D")
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(h)
        style_run(run, size=10, color=WHITE, bold=True)

    for r_idx, row_data in enumerate(rows_data, start=1):
        is_recommended = r_idx == 4
        shade = "F7F7FB" if r_idx % 2 == 1 else "FFFFFF"
        if is_recommended:
            shade = "EEEDFE"
        for c_idx, text in enumerate(row_data):
            cell = table.rows[r_idx].cells[c_idx]
            set_cell_shade(cell, shade)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            style_run(run, size=10, color=DARK_SLATE if is_recommended else CHARCOAL,
                      bold=is_recommended)

    add_paragraph("", space_after=4)


def effort_tag_inline(p, label, color):
    """Small filled rounded tag inline. We approximate using run with colored shading."""
    run = p.add_run(f"  {label}  ")
    style_run(run, size=9, color=WHITE, bold=True)
    rpr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    fill_map = {
        "Decide now": "CC0000",
        "5 min": "22215A",
        "30 min": "F2B95A",
    }
    shd.set(qn("w:fill"), fill_map.get(label, "374151"))
    rpr.append(shd)


def action_row(icon, description, tag_label):
    """Signature 4: 3-column action row with effort tag."""
    table = doc.add_table(rows=1, cols=3)
    remove_table_borders(table)
    table.columns[0].width = Inches(0.4)
    table.columns[1].width = Inches(4.8)
    table.columns[2].width = Inches(1.4)

    icon_cell = table.rows[0].cells[0]
    icon_cell.width = Inches(0.4)
    remove_cell_borders(icon_cell)
    icon_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    icon_p = icon_cell.paragraphs[0]
    icon_p.paragraph_format.space_after = Pt(0)
    ir = icon_p.add_run(icon)
    style_run(ir, size=12, color=DARK_SLATE, bold=True)

    desc_cell = table.rows[0].cells[1]
    desc_cell.width = Inches(4.8)
    remove_cell_borders(desc_cell)
    desc_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    desc_p = desc_cell.paragraphs[0]
    desc_p.paragraph_format.space_after = Pt(0)
    dr = desc_p.add_run(description)
    style_run(dr, color=CHARCOAL)

    tag_cell = table.rows[0].cells[2]
    tag_cell.width = Inches(1.4)
    remove_cell_borders(tag_cell)
    tag_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    tag_p = tag_cell.paragraphs[0]
    tag_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tag_p.paragraph_format.space_after = Pt(0)
    effort_tag_inline(tag_p, tag_label, color=None)

    add_paragraph("", space_after=2)


def summary_callout(body_runs):
    """Signature 5: First-person closing with amber border (left bar + top/bottom)."""
    table = doc.add_table(rows=1, cols=1)
    remove_table_borders(table)
    cell = table.rows[0].cells[0]
    set_cell_shade(cell, AMBER_LIGHT_HEX)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side, sz in (("top", "8"), ("bottom", "8"), ("left", "32")):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:color"), "F2B95A")
        b.set(qn("w:space"), "0")
        tc_borders.append(b)
    tc_pr.append(tc_borders)

    label_p = cell.paragraphs[0]
    label_p.paragraph_format.space_after = Pt(4)
    lr = label_p.add_run("Summary")
    style_run(lr, size=11, color=AMBER_DARK, bold=True)

    body_p = cell.add_paragraph()
    body_p.paragraph_format.space_after = Pt(0)
    for chunk, opts in body_runs:
        run = body_p.add_run(chunk)
        style_run(
            run,
            size=opts.get("size", 11),
            color=opts.get("color", CHARCOAL),
            bold=opts.get("bold", False),
            italic=opts.get("italic", False),
        )

    tc_mar = OxmlElement("w:tcMar")
    for side, val in (("top", "140"), ("bottom", "140"), ("left", "180"), ("right", "180")):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), val)
        m.set(qn("w:type"), "dxa")
        tc_mar.append(m)
    tc_pr.append(tc_mar)

    add_paragraph("", space_after=4)


def calendar_box():
    """Standout calendar item, designed to be selected and pasted into a calendar app."""
    table = doc.add_table(rows=1, cols=1)
    remove_table_borders(table)
    cell = table.rows[0].cells[0]
    set_cell_shade(cell, OFF_WHITE_HEX)
    set_cell_borders(cell, color="1C1C2D", size="6", all_sides=True)

    label_p = cell.paragraphs[0]
    label_p.paragraph_format.space_after = Pt(2)
    lr = label_p.add_run("CALENDAR ITEM")
    style_run(lr, size=10, color=DUSTY_PERIWINKLE, bold=True)

    title_p = cell.add_paragraph()
    title_p.paragraph_format.space_after = Pt(8)
    tr = title_p.add_run("Revisit LingoLinq password manager choice (1Password vs Proton Pass)")
    style_run(tr, size=14, color=DARK_SLATE, bold=True)

    fields = [
        ("Date:", "2027-04-17 (approximately one year from today)"),
        ("Owner:", "Dominic"),
    ]
    for label, value in fields:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        l = p.add_run(f"{label} ")
        style_run(l, color=DARK_SLATE, bold=True)
        v = p.add_run(value)
        style_run(v, color=CHARCOAL)

    ctx_p = cell.add_paragraph()
    ctx_p.paragraph_format.space_before = Pt(6)
    ctx_p.paragraph_format.space_after = Pt(2)
    cl = ctx_p.add_run("Context: ")
    style_run(cl, color=DARK_SLATE, bold=True)
    ct = ctx_p.add_run(
        "Today we committed to 1Password Business. At that time, Proton Pass for Business "
        "was an emerging alternative priced at roughly half the cost ($4.49 vs $7.99 per user) "
        "but too new to trust for our automation. Proton also offers a Business Suite that "
        "bundles Mail, Drive, VPN, and Pass under a single BAA for $12.99/user, which could "
        "eventually replace Google Workspace."
    )
    style_run(ct, color=CHARCOAL)

    check_label = cell.add_paragraph()
    check_label.paragraph_format.space_before = Pt(8)
    check_label.paragraph_format.space_after = Pt(2)
    cr = check_label.add_run("What to check in a year")
    style_run(cr, color=DARK_SLATE, bold=True)

    questions = [
        "Has Proton Pass CLI matured? It launched November 2025, so by April 2027 it will have been in the wild for about 18 months.",
        "Is Proton Business Suite cheaper than our combined Google Workspace plus 1Password spend?",
        "Does Proton offer a free guest seat model comparable to 1Password's 20 included guests?",
        "Has our team grown past 15 people? If so, the per-seat math changes significantly.",
        "Has any client contractually required self-hosted credential management?",
    ]
    for q in questions:
        p = cell.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(q)
        style_run(run, color=CHARCOAL)

    trigger_label = cell.add_paragraph()
    trigger_label.paragraph_format.space_before = Pt(8)
    trigger_label.paragraph_format.space_after = Pt(2)
    tl = trigger_label.add_run("What would trigger an earlier re-evaluation")
    style_run(tl, color=DARK_SLATE, bold=True)

    triggers = [
        "1Password raises prices by more than 20 percent",
        "A 1Password security incident is disclosed",
        "A client signs a contract requiring self-hosted or EU-jurisdiction password management",
        "We grow past 15 full-time team members",
    ]
    for t in triggers:
        p = cell.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(t)
        style_run(run, color=CHARCOAL)

    final = cell.add_paragraph()
    final.paragraph_format.space_before = Pt(6)
    final.paragraph_format.space_after = Pt(0)
    fr = final.add_run("If none of the above, 1Password remains the right call.")
    style_run(fr, color=DARK_SLATE, italic=True)

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, val in (("top", "200"), ("bottom", "200"), ("left", "240"), ("right", "240")):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), val)
        m.set(qn("w:type"), "dxa")
        tc_mar.append(m)
    tc_pr.append(tc_mar)

    add_paragraph("", space_after=4)


def footer_strip():
    footer = doc.sections[0].footer
    footer_table = footer.add_table(rows=1, cols=2, width=Inches(6.6))
    footer_table.autofit = False
    remove_table_borders(footer_table)
    footer_table.columns[0].width = Inches(3.3)
    footer_table.columns[1].width = Inches(3.3)

    left = footer_table.rows[0].cells[0]
    remove_cell_borders(left)
    set_cell_borders(left, color="F0F0F0", size="6", top=True)
    lp = left.paragraphs[0]
    lp.paragraph_format.space_before = Pt(4)
    lp.paragraph_format.space_after = Pt(0)
    lr = lp.add_run("LingoLinq  |  Scot Wahlquist  |  2026")
    style_run(lr, size=9, color=DUSTY_PERIWINKLE)

    right = footer_table.rows[0].cells[1]
    remove_cell_borders(right)
    set_cell_borders(right, color="F0F0F0", size="6", top=True)
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(4)
    rp.paragraph_format.space_after = Pt(0)
    rr = rp.add_run("scot@lingolinq.com")
    style_run(rr, size=9, color=DUSTY_PERIWINKLE)

    default_p = footer.paragraphs[0]
    default_p.text = ""


# ============================================================================
# DOCUMENT CONTENT
# ============================================================================

header_strip()
title_banner()

amber_callout(
    "BOTTOM LINE",
    "I set up 1Password for the company. Trial ends Saturday. It's $23.97/month "
    "for the three of us and free for our long-term contractors. I scanned the full "
    "2026 password manager market and 1Password is still the right call. Need your "
    "approval by Saturday so we don't lose the setup."
)

add_h1("What a password manager actually does for you day-to-day")
add_italic_intro("If you haven't used one before, here's what changes.")
day_to_day = [
    "You install 1Password on your phone and laptop. One app.",
    "When you log in to any website (bank, Stripe, HubSpot, Google), 1Password fills in the username and password for you. You never type a password again.",
    "You have one master password that unlocks the app on your devices. That's the only one you memorize.",
    "When we need to share a login (like the LingoLinq AFCU account), I put it in a shared vault and it shows up on your device automatically. Same with Melissa for dev stuff.",
    "It generates strong passwords for you when you sign up for new sites. No more reusing passwords.",
    "It warns you when a password you use has been leaked in a breach somewhere on the internet.",
]
for line in day_to_day:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.35
    r = p.add_run(line)
    style_run(r)
add_paragraph(
    "The shared-vault piece is the real unlock for us. Right now if I need to give you "
    "access to a new tool, I'm sending passwords in text messages or email, which is "
    "both insecure and a nightmare when passwords change. With 1Password, I update it "
    "once and you see the new password instantly."
)

add_h1("Why a free option doesn't work for us")
add_italic_intro("Bitwarden Free is great for personal use. It is not built for a regulated company.")
add_runs([
    ("There are free password managers, most notably ", {}),
    ("Bitwarden Free", {"bold": True}),
    (". Bitwarden Free is a great product for personal use. I use something like it for my own non-LingoLinq logins.", {}),
])
add_paragraph("The problem: free tiers are built for individuals, not companies. Bitwarden Free does not have:")
free_gaps = [
    ("Shared vaults for teams. ", "You and I literally cannot share a password through the free version. Everyone has their own silo."),
    ("Admin controls. ", "No way for me to say Melissa can see dev secrets but not banking."),
    ("Audit logs. ", "No record of who accessed what password when. Required for compliance."),
    ("A Business Associate Agreement (BAA). ", "This is the legal document HIPAA requires from any vendor that touches healthcare data indirectly. Free tools do not sign BAAs."),
]
for lead, tail in free_gaps:
    add_bullet(tail, bold_lead=lead)
add_runs([
    ("So we're not really comparing 1Password Business vs free. We're comparing 1Password Business vs the ", {}),
    ("paid business tier of Bitwarden", {"bold": True}),
    (", because that's the only version of Bitwarden with the features a real company needs.", {}),
])

add_h1("Why this matters for LingoLinq specifically")
add_italic_intro("Procurement teams have checklists. We need to check the boxes or we lose deals.")
add_paragraph("We sell to schools, hospitals, and European clients. Every one of those buyers sends us a security questionnaire before they sign a contract. Questions like:")
for q in [
    "How do you manage administrative credentials?",
    "Do you have audit logs of who accesses production systems?",
    "Can you provide a BAA?",
    "How do you revoke access when someone leaves the team?",
]:
    add_bullet(q)
add_paragraph(
    "If we answer with a shared Google Doc or free Bitwarden, we lose deals. Not "
    "because it's insecure in practice (though it is), but because district IT and "
    "hospital compliance teams have checklists. This is a competitive issue, not "
    "just a risk issue."
)
add_paragraph("We're also legally on the hook under:")
frameworks = [
    ("HIPAA ", "(hospital clients process patient data)"),
    ("FERPA ", "(schools process student education records)"),
    ("COPPA ", "(most of our end users are children under 13)"),
    ("GDPR ", "(European clients have EU data protection rules)"),
]
for lead, tail in frameworks:
    add_bullet(tail, bold_lead=lead)
add_paragraph(
    "These frameworks all require controls around who accesses sensitive systems and "
    "proof of those controls. A real password manager with audit logs and individual "
    "accounts is how we provide that proof."
)

add_h1("The honest cost comparison")
add_italic_intro("On sticker price Bitwarden wins by $72/year. The picture changes when you factor in guest seats and migration cost.")
cost_table()
add_runs([
    ("The real comparison is $18/mo vs $24/mo. ", {}),
    ("Bitwarden Enterprise would save us $72 per year.", {"bold": True}),
])

add_h1("Why I'm still recommending 1Password over Bitwarden")
add_italic_intro("Three reasons that flip the math.")

add_h2("1. Free contractor access")
add_paragraph(
    "1Password Business includes 20 free guest seats. That covers Traci, Brian "
    "(OpenAAC, joining in a few weeks), and any AI interns or design contractors "
    "we bring on through the rest of this year."
)
add_runs([
    ("Bitwarden charges for every user. If we add 3 contractors over the next 12 months, Bitwarden becomes ", {}),
    ("more expensive", {"bold": True}),
    (" than 1Password. That $72/year savings flips to a net loss.", {}),
])

add_h2("2. We already built automation on top of 1Password")
add_paragraph("I've been setting this up for a few weeks. What's already working:")
for line in [
    "Server secrets (database passwords, API keys) auto-sync from 1Password to Render every hour via a GitHub Action. This fixed a real production bug last month where workers had stale keys.",
    "Claude Code and my other AI tools pull credentials from 1Password instead of having them sit in .env files on my laptop.",
    "Vault structure, access rules, and service accounts are all configured.",
]:
    add_bullet(line)
add_runs([
    ("Switching to Bitwarden means Melissa rebuilds all of that. Conservatively 2 to 3 days of her contractor time. At her rate, that's ", {}),
    ("more than a year of 1Password subscription cost", {"bold": True}),
    (" out the window in the first week.", {}),
])

add_h2("3. Audit-ready today")
add_paragraph(
    "If a district asks us tomorrow to prove we manage credentials properly, "
    "1Password Business gives us the audit log, BAA, and SOC 2 report in one click. "
    "We don't have to scramble."
)

add_h1("Where Bitwarden would genuinely win")
add_italic_intro("I want to be fair. Here's when I'd recommend the other way.")
add_paragraph("Bitwarden would be the better choice if:")
for line in [
    "We were starting from scratch with no automation built (we're not)",
    "We had 15 or more employees where the per-seat math compounds (we have 3)",
    "We contractually needed to self-host our password manager (no client has asked)",
    "Open-source software was a company principle we market on (it's not)",
]:
    add_bullet(line)
add_paragraph(
    "None of those apply right now. If we grow to 15 people or land a client that "
    "requires self-hosted infrastructure, we revisit. For now, 1Password is the right call."
)

add_h1("Action items")
add_italic_intro("Three things from you. The first one is the one with the deadline.")
action_row("🔴", "Approve the $23.97/month charge before Saturday so we don't lose the trial setup", "Decide now")
action_row("📝", "Accept your pending 1Password invite (check email from 1Password). I'll add you to the Co-Founders vault once you accept.", "5 min")
action_row("🔐", "Tomorrow's meeting: rotate the AFCU password and set up proper 2FA that works for both of us remotely", "30 min")
add_paragraph(
    "If you want to talk any of this through before Saturday, grab me on Chat or "
    "we can hop on a call."
)

add_h1("One year from now")
add_italic_intro("Drop this into your task calendar so we don't forget to revisit.")
calendar_box()

add_h1("Appendix: Market scan as of 2026-04-17")
add_italic_intro("I checked the full landscape before signing. Here's what I found.")
add_paragraph(
    "Before finalizing this decision, I ran a fresh scan of the business password "
    "manager market to make sure no new player had emerged that would be a better "
    "fit for us. Summary of findings:"
)
market_notes = [
    ("Proton Pass Business ", "is the only genuinely new contender worth naming (CLI launched Nov 2025). Cheaper but with less mature automation. Revisit in 12 months."),
    ("Keeper Business ", "is comparable to 1Password but charges per guest seat."),
    ("Dashlane ", "killed its free tier in Sep 2025, no publicly confirmed BAA."),
    ("NordPass ", "has SOC 2 but does not publicly advertise HIPAA BAA support."),
    ("LastPass ", "ruled out permanently (2022 breach still causing damage, ICO penalty Nov 2025, unpatched DEF CON 33 vuln from Aug 2025)."),
    ("Passbolt and Psono ", "are open-source self-hosted options. Not worth the operational burden for a 3-person team."),
]
for lead, tail in market_notes:
    add_bullet(tail, bold_lead=lead)
add_runs([
    ("A February 2026 independent study from ETH Zurich and USI tested 27 theoretical attack scenarios against major password managers. ", {}),
    ("1Password had the fewest findings (2). Bitwarden had 12.", {"bold": True}),
    (" That matters for procurement questionnaires.", {}),
])
add_paragraph(
    "No breaches were disclosed for 1Password, Bitwarden, Proton Pass, Keeper, or "
    "Dashlane during the 2023 to 2026 window."
)

summary_callout([
    ("I am confident this is the right call for where we are right now. ", {"color": CHARCOAL}),
    ("The $24/month is the cheapest insurance we can buy against a procurement disqualification, "
     "and I've already built the automation we need on top of it. If you have any hesitation, "
     "the calendar item above gives us a clean checkpoint to revisit in a year.", {"color": CHARCOAL}),
])

footer_strip()

output_path = "/home/scotw/MEMO_Dominic_1Password_Decision.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
